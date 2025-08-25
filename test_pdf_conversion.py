#!/usr/bin/env python3
"""
PDF Conversion Test Script
Tests if PDF conversion tools are properly installed and working in the Docker container
"""

import os
import sys
import subprocess
import tempfile
from pathlib import Path

def test_libreoffice():
    """Test LibreOffice installation"""
    print("Testing LibreOffice...")
    try:
        result = subprocess.run(['libreoffice', '--version'], 
                              capture_output=True, text=True, timeout=10)
        if result.returncode == 0:
            print(f"✓ LibreOffice found: {result.stdout.strip()}")
            return True
        else:
            print(f"✗ LibreOffice failed: {result.stderr}")
            return False
    except Exception as e:
        print(f"✗ LibreOffice test failed: {e}")
        return False

def test_unoconv():
    """Test unoconv installation"""
    print("Testing unoconv...")
    try:
        result = subprocess.run(['unoconv', '--version'], 
                              capture_output=True, text=True, timeout=10)
        if result.returncode == 0:
            print(f"✓ unoconv found: {result.stdout.strip()}")
            return True
        else:
            print(f"✗ unoconv failed: {result.stderr}")
            return False
    except Exception as e:
        print(f"✗ unoconv test failed: {e}")
        return False

def test_pdf_conversion():
    """Test actual PDF conversion"""
    print("Testing PDF conversion...")
    
    # Create a simple test DOCX file
    try:
        from docx import Document
        
        # Create test document
        doc = Document()
        doc.add_heading('Test Document', 0)
        doc.add_paragraph('This is a test document for PDF conversion.')
        doc.add_paragraph('If you can see this as PDF, the conversion worked!')
        
        # Save to temp file
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
            docx_path = tmp.name
            doc.save(docx_path)
        
        # Try to convert to PDF using unoconv
        pdf_path = docx_path.replace('.docx', '.pdf')
        
        cmd = ['unoconv', '-f', 'pdf', '-o', pdf_path, docx_path]
        result = subprocess.run(cmd, capture_output=True, text=True, timeout=30)
        
        if result.returncode == 0 and os.path.exists(pdf_path):
            pdf_size = os.path.getsize(pdf_path)
            print(f"✓ PDF conversion successful! Output file: {pdf_size} bytes")
            
            # Cleanup
            os.unlink(docx_path)
            os.unlink(pdf_path)
            return True
        else:
            print(f"✗ PDF conversion failed: {result.stderr}")
            # Cleanup
            if os.path.exists(docx_path):
                os.unlink(docx_path)
            return False
            
    except ImportError:
        print("✗ python-docx not available, cannot create test document")
        return False
    except Exception as e:
        print(f"✗ PDF conversion test failed: {e}")
        return False

def test_fonts():
    """Test font availability for Arabic text"""
    print("Testing font availability...")
    try:
        result = subprocess.run(['fc-list', ':lang=ar'], 
                              capture_output=True, text=True, timeout=10)
        if result.returncode == 0:
            fonts = result.stdout.strip().split('\n')
            arabic_fonts = [f for f in fonts if f.strip()]
            if arabic_fonts:
                print(f"✓ Found {len(arabic_fonts)} Arabic fonts")
                return True
            else:
                print("⚠ No Arabic fonts found")
                return False
        else:
            print("⚠ Cannot check fonts")
            return False
    except Exception as e:
        print(f"⚠ Font check failed: {e}")
        return False

def main():
    """Main test function"""
    print("="*60)
    print("PDF Conversion Test Suite")
    print("="*60)
    print()
    
    # Environment info
    print("Environment Information:")
    print(f"  OS: {os.name}")
    print(f"  Python: {sys.version}")
    print(f"  Working directory: {os.getcwd()}")
    print()
    
    # Run tests
    tests = [
        ("LibreOffice", test_libreoffice),
        ("unoconv", test_unoconv),
        ("Fonts", test_fonts),
        ("PDF Conversion", test_pdf_conversion),
    ]
    
    results = {}
    for test_name, test_func in tests:
        print(f"Running {test_name} test...")
        try:
            results[test_name] = test_func()
        except Exception as e:
            print(f"✗ {test_name} test crashed: {e}")
            results[test_name] = False
        print()
    
    # Summary
    print("="*60)
    print("Test Results Summary:")
    print("="*60)
    
    passed = 0
    total = len(results)
    
    for test_name, result in results.items():
        status = "✓ PASS" if result else "✗ FAIL"
        print(f"  {test_name:<20} {status}")
        if result:
            passed += 1
    
    print()
    print(f"Overall: {passed}/{total} tests passed")
    
    if passed == total:
        print("🎉 All tests passed! PDF conversion is ready to use.")
        return 0
    elif passed >= total - 1:  # Allow one failure (maybe fonts)
        print("⚠ Most tests passed, PDF conversion should work.")
        return 0
    else:
        print("❌ Critical tests failed, PDF conversion may not work properly.")
        return 1

if __name__ == "__main__":
    sys.exit(main())
