"""
Enhanced Arabic Worksheet Generator
A comprehensive system for generating professional Arabic worksheets with smart layout adaptation.
Consolidates functionality from multiple files into a clean, maintainable structure.
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.shared import OxmlElement, qn
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from abc import ABC, abstractmethod
from typing import Dict, List, Any, Optional, Union
from itertools import cycle
import json
import datetime
from dataclasses import dataclass, field
from enum import Enum
import re

# ============================================================================
# CONFIGURATION AND DATA MODELS
# ============================================================================

class LayoutType(Enum):
    """Defines available layout types"""
    FULL_WIDTH = "full_width"
    WITH_SIDEBAR = "with_sidebar"

# ----------------------------------------------------------------------------
# GLOBAL CONSTANTS AND UI CONFIGURATION
# ----------------------------------------------------------------------------

# Word uses DXA units (twips/20th of a point) for widths: 1 inch = 1440 dxa
DXA_PER_INCH: int = 1440

@dataclass
class DimensionsConfig:
    """All dimensional constants in a single place (inches/dxa)."""
    page_width_in: float = 7.9
    top_margin_in: float = 0.4
    bottom_margin_in: float = 0.4
    left_margin_in: float = 0.3
    right_margin_in: float = 0.3

    # Two-column layout widths
    main_width_in: float = 6.0
    sidebar_width_in: float = 1.9

    # Header widths
    header_cell_width_in: float = 2.0
    name_class_cell_width_in: float = 4.0

    # Sidebar box width
    sidebar_box_width_in: float = 1.2

    # Table cell internal margins (DXA)
    cell_margin_dxa: int = 100

@dataclass
class FontSizeConfig:
    """Font sizes for various parts (in points)."""
    header_dark_pt: int = 11
    header_light_pt: int = 12
    title_pt: int = 16
    section_header_pt: int = 13
    mc_number_pt: int = 12
    mc_text_pt: int = 12
    choice_label_pt: int = 10
    choice_text_pt: int = 10
    essay_answer_pt: int = 10
    level_text_pt: int = 11
    footer_pt: int = 9

@dataclass
class BorderConfig:
    """Default border style settings."""
    width_pt: int = 2
    style: str = 'dashed'
    # Word XML expects size in eighths of a point (approx); often sz ~ width*5 is used
    sz_multiplier: int = 5

@dataclass
class ChoiceGridConfig:
    """Choice grid configuration (kept 2x2, but centralized)."""
    rows: int = 2
    cols: int = 2

@dataclass
class EssayConfig:
    """Essay rendering configuration."""
    answer_lines: int = 5

@dataclass
class HeaderConfig:
    """Header configuration data"""
    subject_memo: str = "مذكرة الرياضيات"
    worksheet_number: str = "ورقة عمل (2)"
    name_label: str = "الأسم:"
    class_label: str = "الصف:"
    semester: str = "الفصل المتوسط"
    grade: str = "ثالث متوسط"

@dataclass
class ChoiceLabels:
    """Choice labeling configuration"""
    symbols: List[str] = field(default_factory=lambda: ['أ', 'ب', 'ج', 'د'])
    numbers: List[str] = field(default_factory=lambda: ['١', '٢', '٣', '٤'])
    separator: str = " - "
    numbering_type: str = "numbers"  # "symbols" or "numbers"

@dataclass
class SidebarLabels:
    """Sidebar section labels"""
    before_lesson_title: str = "ما قبل الدرس"
    goals_title: str = "أهداف الدرس"
    applications_title: str = "تطبيقات الدرس"
    levels_title: str = "مستويات الاتقان"
    notes_title: str = "ملاحظات المعلم"
    goals_bullet: str = "✓ "
    levels_bullet: str = "〇 \t"

@dataclass
class GeneralLabels:
    """General text labels"""
    section_icon: str = "⚙ "
    answer_prefix: str = "الإجابة: "
    underline_char: str = "_"
    underline_count: int = 100

@dataclass
class UIConfig:
    """Bundles all UI-related configurations to avoid magic numbers."""
    dimensions: DimensionsConfig = field(default_factory=DimensionsConfig)
    fonts: FontSizeConfig = field(default_factory=FontSizeConfig)
    borders: BorderConfig = field(default_factory=BorderConfig)
    choice_grid: ChoiceGridConfig = field(default_factory=ChoiceGridConfig)
    essay: EssayConfig = field(default_factory=EssayConfig)

@dataclass
class Colors:
    """Color scheme configuration"""
    dark_navy: RGBColor = field(default_factory=lambda: RGBColor(52, 63, 85))
    medium_blue: RGBColor = field(default_factory=lambda: RGBColor(78, 96, 128))
    light_blue: RGBColor = field(default_factory=lambda: RGBColor(149, 165, 195))
    yellow_highlight: RGBColor = field(default_factory=lambda: RGBColor(255, 255, 102))
    green_level: RGBColor = field(default_factory=lambda: RGBColor(144, 238, 144))
    orange_level: RGBColor = field(default_factory=lambda: RGBColor(255, 200, 124))
    red_level: RGBColor = field(default_factory=lambda: RGBColor(255, 160, 160))
    light_gray: RGBColor = field(default_factory=lambda: RGBColor(245, 245, 245))
    border_gray: RGBColor = field(default_factory=lambda: RGBColor(180, 180, 180))
    white: RGBColor = field(default_factory=lambda: RGBColor(255, 255, 255))
    black: RGBColor = field(default_factory=lambda: RGBColor(0, 0, 0))
    text_blue: RGBColor = field(default_factory=lambda: RGBColor(40, 55, 80))

@dataclass
class Question:
    """Base question data structure"""
    question: str
    number: Optional[int] = None

@dataclass
class MultipleChoiceQuestion(Question):
    """Multiple choice question with validation"""
    choices: List[str] = field(default_factory=list)
    answer_key: Optional[int] = None
    
    def __post_init__(self):
        # if self.choices and len(self.choices) != 4:
        #     raise ValueError("Multiple choice questions must have exactly 4 choices")
        if self.answer_key is not None and not (0 <= self.answer_key < 4):
            raise ValueError("Answer key must be between 0 and 3")

@dataclass
class EssayQuestion(Question):
    """Essay question with optional answer"""
    answer: str = ""

@dataclass
class SidebarConfig:
    """Sidebar configuration with defaults"""
    before_lesson: str = ""
    goals: List[str] = field(default_factory=list)
    applications: List[str] = field(default_factory=list)
    levels: List[str] = field(default_factory=list)
    notice: str = ""

@dataclass
class WorksheetData:
    """Complete worksheet data structure"""
    title: str = ""
    multiple_choice_header: str = ""
    essay_header: str = ""
    multiple_choice_questions: List[MultipleChoiceQuestion] = field(default_factory=list)
    essay_questions: List[EssayQuestion] = field(default_factory=list)
    sidebar: Optional[SidebarConfig] = None
    layout_type: LayoutType = LayoutType.FULL_WIDTH
    header_config: HeaderConfig = field(default_factory=HeaderConfig)
    choice_labels: ChoiceLabels = field(default_factory=ChoiceLabels)
    sidebar_labels: SidebarLabels = field(default_factory=SidebarLabels)
    general_labels: GeneralLabels = field(default_factory=GeneralLabels)
    
    def __post_init__(self):
        # Determine layout based on sidebar presence
        if self.sidebar is not None:
            self.layout_type = LayoutType.WITH_SIDEBAR
        
        # Validate at least one question exists
        if not self.multiple_choice_questions and not self.essay_questions:
            raise ValueError("At least one question must be provided")

# ============================================================================
# FORMATTING AND RENDERING SYSTEM
# ============================================================================

# ----------------------------------------------------------------------------------------------------------------------
# low-level paragraph formatter (keeps name)
def format_paragraph_rtl(paragraph, preserve_alignment: bool = False) -> None:
    """Apply RTL formatting to a paragraph while optionally preserving alignment"""
    # Ensure paragraph property container exists
    pPr = paragraph._element.get_or_add_pPr()

    # Set paragraph-level bidi (logical right-to-left)
    # <w:bidi w:val="1"/> - indicates paragraph is RTL
    bidi = OxmlElement('w:bidi')
    # Some Word readers accept empty <w:bidi/>; set val to '1' for clarity
    try:
        bidi.set(qn('w:val'), '1')
    except Exception:
        # some python-docx versions ignore attribute; ignore silently
        pass
    # If a bidi element already present, remove duplicates
    # (safe-append)
    pPr.append(bidi)

    # Remove spacing around paragraph (keep behaviour)
    paragraph_format = paragraph.paragraph_format
    paragraph_format.space_after = Pt(0)
    paragraph_format.space_before = Pt(0)

    # For RTL paragraphs, default alignment should be RIGHT unless preserve_alignment == True
    if not preserve_alignment:
        paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

# ----------------------------------------------------------------------------------------------------------------------
# Helper functions used inside the class (kept internal)
def contains_arabic(text: str) -> bool:
    """Return True if text contains Arabic/Hebrew characters (broadly)."""
    if not text:
        return False
    # Arabic, Arabic supplement, Arabic extended ranges and Hebrew
    return bool(re.search(
        r'[\u0590-\u08FF\uFB1D-\uFDFF\uFE70-\uFEFF]',
        text
    ))

def contains_strong_latin(text: str) -> bool:
    """Return True if text contains primarily Latin letters or digits."""
    if not text:
        return False
    return bool(re.search(r'[A-Za-z0-9]', text))

def run_contains_omml(run) -> bool:
    """Detect if a run contains OMML (Word Math) elements in its XML."""
    # look for 'oMath' localname in any descendant tag
    for child in run._element.iter():
        tag = getattr(child, 'tag', '')
        if isinstance(tag, str) and tag.endswith('oMath'):
            return True
    return False

def set_run_rtl(run, on: bool = True) -> None:
    """Set run-level RTL property: <w:rPr><w:rtl w:val="1"/></w:rPr>"""
    try:
        rPr = run._r.get_or_add_rPr()
    except AttributeError:
        # fallback: append rPr element explicitly
        rPr = OxmlElement('w:rPr')
        run._r.append(rPr)
    rtl = OxmlElement('w:rtl')
    rtl.set(qn('w:val'), '1' if on else '0')
    rPr.append(rtl)

# Unicode directional controls & punctuation map
RLM = '\u200F'       # Right-to-Left Mark
LRM = '\u200E'       # Left-to-Right Mark
LRI = '\u2066'       # Left-to-Right Isolate
RLI = '\u2067'       # Right-to-Left Isolate
PDI = '\u2069'       # Pop Directional Isolate

ARABIC_PUNCT_MAP = {',': '\u060C',  # Arabic comma
                    '?': '\u061F',  # Arabic question mark
                    ';': '\u061B'}  # Arabic semicolon

# ----------------------------------------------------------------------------------------------------------------------
# ArabicTextFormatter class (names preserved)
class ArabicTextFormatter:
    """Handles Arabic text formatting and bidirectional text"""

    @staticmethod
    def detect_text_direction(text: str) -> str:
        """Detect if text is primarily RTL, LTR, or mixed."""
        if not text:
            return "neutral"

        arabic_count = 0
        latin_count = 0
        for ch in text:
            o = ord(ch)
            # Arabic & Hebrew blocks and presentation forms
            if (0x0590 <= o <= 0x08FF) or (0xFB1D <= o <= 0xFEFC):
                arabic_count += 1
            # Latin / ASCII letters & digits
            elif (0x0041 <= o <= 0x007A) or (0x0030 <= o <= 0x0039):
                latin_count += 1

        total = arabic_count + latin_count
        if total == 0:
            return "neutral"
        arabic_ratio = arabic_count / total
        if arabic_ratio > 0.6:
            return "rtl"
        elif arabic_ratio < 0.2:
            return "ltr"
        else:
            return "mixed"

    @staticmethod
    def format(text: str) -> str:
        """Return text as-is (Word handles shaping). Kept for API stability."""
        return text

    @staticmethod
    def apply_smart_rtl_to_paragraph(paragraph, text: Optional[str] = None, preserve_alignment: bool = False) -> None:
        """
        Apply RTL formatting intelligently based on text content.
        - Detects paragraph direction (rtl/ltr/mixed)
        - Sets paragraph bidi where needed
        - Replaces punctuation with Arabic equivalents when paragraph is Arabic
        - Wraps LTR/mathed runs using isolates (LRI/RLI + PDI)
        - Anchors punctuation using RLM or LRM heuristics
        """
        # collect paragraph text if not provided
        if text is None:
            text = ''.join([r.text for r in paragraph.runs])

        direction = ArabicTextFormatter.detect_text_direction(text)

        if direction == "rtl":
            # make paragraph RTL
            format_paragraph_rtl(paragraph, preserve_alignment)

            # convert punctuation and wrap LTR runs
            for run in paragraph.runs:
                # replace punctuation in Arabic paragraphs for neutral punctuation
                if contains_arabic(run.text):
                    for k, v in ARABIC_PUNCT_MAP.items():
                        if k in run.text:
                            run.text = run.text.replace(k, v)

                # If run contains LTR strong content or math, isolate it
                if contains_strong_latin(run.text) or run_contains_omml(run):
                    # wrap with LRI...PDI (safer than embedding)
                    run.text = f"{LRI}{run.text}{PDI}"

                    # If it's OMML math, append RLM to anchor following punctuation visually to RTL sentence
                    if run_contains_omml(run):
                        run.text = run.text + RLM

                # If run is purely Arabic letters, ensure run-level RTL property for stronger rendering
                if contains_arabic(run.text) and not contains_strong_latin(run.text):
                    try:
                        set_run_rtl(run, on=True)
                    except Exception:
                        # If run-level setting fails, we still keep paragraph level RTL
                        pass

        elif direction == "mixed":
            # Apply paragraph-level RTL but handle isolates and punctuation carefully
            format_paragraph_rtl(paragraph, preserve_alignment)

            for run in paragraph.runs:
                text_run = run.text

                # replace punctuation for Arabic-dominant runs
                if contains_arabic(text_run) and not contains_strong_latin(text_run):
                    for k, v in ARABIC_PUNCT_MAP.items():
                        text_run = text_run.replace(k, v)

                # For runs that look LTR, wrap in LRI...PDI
                if contains_strong_latin(text_run) or run_contains_omml(run):
                    text_run = f"{LRI}{text_run}{PDI}"
                    if run_contains_omml(run):
                        text_run = text_run + RLM

                # apply run-level RTL if Arabic
                if contains_arabic(text_run) and not contains_strong_latin(text_run):
                    try:
                        set_run_rtl(run, on=True)
                    except Exception:
                        pass

                run.text = text_run

            # add an explicit logical bidi override element (w:bidiVisual = 0)
            pPr = paragraph._element.get_or_add_pPr()
            bidiOverride = OxmlElement('w:bidiVisual')
            try:
                bidiOverride.set(qn('w:val'), '0')
            except Exception:
                pass
            pPr.append(bidiOverride)

        elif direction == "ltr":
            # Paragraph is LTR: leave it mostly untouched, but normalize spacing
            para_fmt = paragraph.paragraph_format
            para_fmt.space_after = Pt(0)
            para_fmt.space_before = Pt(0)
            if not preserve_alignment:
                para_fmt.alignment = WD_ALIGN_PARAGRAPH.LEFT

            # But for Arabic runs inside LTR paragraphs, add RLI isolates for the Arabic fragments
            for run in paragraph.runs:
                if contains_arabic(run.text) and not contains_strong_latin(run.text):
                    run.text = f"{RLI}{run.text}{PDI}"
                    # force run-level RTL for Arabic chunk
                    try:
                        set_run_rtl(run, on=True)
                    except Exception:
                        pass

        else:
            # Neutral: minimal formatting but keep spacing clean
            para_fmt = paragraph.paragraph_format
            para_fmt.space_after = Pt(0)
            para_fmt.space_before = Pt(0)

    @staticmethod
    def apply_rtl_to_paragraph(paragraph, preserve_alignment: bool = False) -> None:
        """Apply RTL formatting to a paragraph (legacy/explicit method)"""
        format_paragraph_rtl(paragraph, preserve_alignment)

class DocumentStyler:
    """Handles all document styling operations"""
    
    def __init__(self, colors: Colors, ui: UIConfig):
        self.colors = colors
        self.ui = ui
    
    def setup_document(self, doc: Document) -> None:
        """Configure document margins and settings"""
        dims = self.ui.dimensions
        for section in doc.sections:
            section.top_margin = Inches(dims.top_margin_in)
            section.bottom_margin = Inches(dims.bottom_margin_in)
            section.left_margin = Inches(dims.left_margin_in)
            section.right_margin = Inches(dims.right_margin_in)
    
    def set_cell_background(self, cell, color: RGBColor) -> None:
        """Set cell background color"""
        cell_xml = cell._tc
        tcPr = cell_xml.get_or_add_tcPr()
        shade_obj = OxmlElement('w:shd')
        shade_obj.set(qn('w:fill'), f"{color[0]:02x}{color[1]:02x}{color[2]:02x}")
        tcPr.append(shade_obj)
        
    
    def set_cell_border(self, cell, color: RGBColor = None, width: Optional[int] = None) -> None:
        """Add border to cell"""
        cell_xml = cell._tc
        tcPr = cell_xml.get_or_add_tcPr()
        tcBorders = OxmlElement('w:tcBorders')
        
        border_color = color or self.colors.dark_navy
        border_width = width if width is not None else self.ui.borders.width_pt
        border_sz = str(border_width * self.ui.borders.sz_multiplier)
        border_style = self.ui.borders.style
        
        for border_name in ['top', 'left', 'bottom', 'right']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), border_style)
            border.set(qn('w:sz'), border_sz)
            border.set(qn('w:color'), f"{border_color[0]:02x}{border_color[1]:02x}{border_color[2]:02x}")
            tcBorders.append(border)
        
        tcPr.append(tcBorders)
    
    def remove_table_borders(self, table) -> None:
        """Remove default table borders"""
        for row in table.rows:
            for cell in row.cells:
                cell._element.get_or_add_tcPr().append(
                    parse_xml(r'<w:tcBorders %s><w:top w:val="nil"/><w:left w:val="nil"/><w:bottom w:val="nil"/><w:right w:val="nil"/></w:tcBorders>' % nsdecls('w'))
                )

class LayoutManager:
    """Manages document layout creation"""
    
    def __init__(self, styler: DocumentStyler):
        self.styler = styler
    
    def create_layout(self, doc: Document, layout_type: LayoutType) -> tuple:
        """Create document layout based on type"""
        if layout_type == LayoutType.WITH_SIDEBAR:
            return self._create_sidebar_layout(doc)
        else:
            return self._create_full_width_layout(doc)
    
    def _create_sidebar_layout(self, doc: Document) -> tuple:
        """Create two-column layout with sidebar"""
        table = doc.add_table(rows=1, cols=2)
        table.autofit = True
        table.width = Inches(self.styler.ui.dimensions.page_width_in)
        
        # Set column widths
        
        main_cell = table.cell(0, 0)
        sidebar_cell = table.cell(0, 1)
        
        # Apply XML width constraints
        dims = self.styler.ui.dimensions
        for cell, width in [(main_cell, dims.main_width_in), (sidebar_cell, dims.sidebar_width_in)]:
            cell._tc.get_or_add_tcPr().append(
                parse_xml(f'<w:tcW w:w="{int(width * DXA_PER_INCH)}" w:type="dxa" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')
            )
        
        self._set_cell_margins(main_cell, sidebar_cell)
        self.styler.remove_table_borders(table)
        
        return main_cell, sidebar_cell
    
    def _create_full_width_layout(self, doc: Document) -> tuple:
        """Create single-column full-width layout"""
        table = doc.add_table(rows=1, cols=1)
        table.autofit = True
        table.width = Inches(self.styler.ui.dimensions.page_width_in)
        
        main_cell = table.cell(0, 0)
        
        # Set full width
        main_cell._tc.get_or_add_tcPr().append(
            parse_xml(f'<w:tcW w:w="{int(self.styler.ui.dimensions.page_width_in * DXA_PER_INCH)}" w:type="dxa" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')
        )
        
        self._set_cell_margins(main_cell)
        self.styler.remove_table_borders(table)
        
        return main_cell, None
    
    def _set_cell_margins(self, *cells) -> None:
        """Set cell margins for proper spacing"""
        margin = self.styler.ui.dimensions.cell_margin_dxa
        for cell in cells:
            cell._tc.get_or_add_tcPr().append(
                parse_xml(
                    (
                        '<w:tcMar xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                        f'<w:top w:w="{margin}" w:type="dxa"/>'
                        f'<w:left w:w="{margin}" w:type="dxa"/>'
                        f'<w:bottom w:w="{margin}" w:type="dxa"/>'
                        f'<w:right w:w="{margin}" w:type="dxa"/>'
                        '</w:tcMar>'
                    )
                )
            )

# ============================================================================
# CONTENT RENDERERS
# ============================================================================

class ContentRenderer(ABC):
    """Abstract base for content renderers"""
    
    def __init__(self, styler: DocumentStyler, formatter: ArabicTextFormatter):
        self.styler = styler
        self.formatter = formatter
    
    @abstractmethod
    def render(self, container, data: Any) -> None:
        pass

class HeaderRenderer(ContentRenderer):
    """Renders document headers"""

    def render(self, doc: Document, title: str = None, header_config: HeaderConfig = None) -> None:
        """Render document header with title"""
        if header_config is None:
            header_config = HeaderConfig()
        self._render_info_header(doc, header_config)
        if title:
            self._render_title(doc, title)

    def _render_info_header(self, doc: Document, header_config: HeaderConfig) -> None:
        """Render the information header section"""
        table = doc.add_table(rows=2, cols=3)
        table.autofit = True
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Configure header content from config
        header_data = [
            [(header_config.subject_memo, True), (header_config.name_label, False), (header_config.semester, True)],
            [(header_config.worksheet_number, True), (header_config.class_label, False), (header_config.grade, True)]
        ]

        for row_idx, row_data in enumerate(header_data):
            for col_idx, (text, is_dark) in enumerate(row_data):
                cell = table.cell(row_idx, col_idx)
                cell.width = Inches(self.styler.ui.dimensions.header_cell_width_in)

                para = cell.paragraphs[0]
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = para.add_run(self.formatter.format(text))
                run.font.size = Pt(self.styler.ui.fonts.header_dark_pt if is_dark else self.styler.ui.fonts.header_light_pt)
                run.bold = True

                # Apply RTL while preserving CENTER alignment
                self.formatter.apply_rtl_to_paragraph(para, preserve_alignment=True)

                if is_dark:
                    run.font.color.rgb = self.styler.colors.white
                    self.styler.set_cell_background(cell, self.styler.colors.dark_navy)
                else:
                    run.font.color.rgb = self.styler.colors.text_blue
                    self.styler.set_cell_background(cell, self.styler.colors.light_gray)

                if col_idx == 1:
                    self.styler.set_cell_border(cell, self.styler.colors.border_gray)
                    cell._tc.get_or_add_tcPr().append(
                        parse_xml(f'<w:tcW w:w="{int(self.styler.ui.dimensions.name_class_cell_width_in * DXA_PER_INCH)}" w:type="dxa" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')
                    )
                    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    # Re-apply RTL without preserving alignment for the name/class fields
                    self.formatter.apply_rtl_to_paragraph(para, preserve_alignment=False)
        doc.add_paragraph()

    def _render_title(self, doc: Document, title: str) -> None:
        """Render document title"""
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(self.formatter.format(title))
        run.font.size = Pt(self.styler.ui.fonts.title_pt)
        run.bold = True
        run.font.color.rgb = self.styler.colors.dark_navy

        # Apply smart RTL while preserving CENTER alignment
        self.formatter.apply_smart_rtl_to_paragraph(para, title, preserve_alignment=True)
        

class SectionHeaderRenderer(ContentRenderer):
    """Renders section headers with instructions"""

    def render(self, container, instruction: str, general_labels: GeneralLabels = None) -> None:
        """Render instruction section header"""
        if general_labels is None:
            general_labels = GeneralLabels()

        table = container.add_table(rows=1, cols=1)
        table.autofit = True
        table.alignment = WD_TABLE_ALIGNMENT.RIGHT

        cell = table.cell(0, 0)
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run = para.add_run(general_labels.section_icon + self.formatter.format(instruction))
        run.font.size = Pt(self.styler.ui.fonts.section_header_pt)
        run.bold = True
        run.font.color.rgb = self.styler.colors.white

        # Apply smart RTL while preserving CENTER alignment
        self.formatter.apply_smart_rtl_to_paragraph(para, instruction, preserve_alignment=True)

        self.styler.set_cell_background(cell, self.styler.colors.dark_navy)

class MultipleChoiceRenderer(ContentRenderer):
    """Renders multiple choice questions"""

    def render(self, container, question: MultipleChoiceQuestion, choice_labels: ChoiceLabels = None) -> None:
        """Render a multiple choice question"""
        if choice_labels is None:
            choice_labels = ChoiceLabels()

        # Main question container
        q_table = container.add_table(rows=1, cols=1)
        q_table.autofit = True
        q_table.alignment = WD_TABLE_ALIGNMENT.LEFT
        q_cell = q_table.cell(0, 0)
        q_cell._tc.get_or_add_tcPr().append(
            parse_xml(f'<w:tcW w:w="{int(self.styler.ui.dimensions.page_width_in * DXA_PER_INCH)}" w:type="dxa" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')
        )

        # Question header
        self._render_question_header(q_cell, question)

        # Choices grid
        self._render_choices_grid(q_cell, question, choice_labels)

        # Styling
        self.styler.set_cell_border(q_cell, self.styler.colors.dark_navy, self.styler.ui.borders.width_pt)

    def _render_question_header(self, cell, question: MultipleChoiceQuestion) -> None:
        """Render question header with number"""
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Question number
        num_run = para.add_run(f" {question.number} ")
        num_run.font.size = Pt(self.styler.ui.fonts.mc_number_pt)
        num_run.bold = True
        num_run.font.color.rgb = self.styler.colors.white

        # Question text
        text_run = para.add_run(f"  {self.formatter.format(question.question)}")
        text_run.font.size = Pt(self.styler.ui.fonts.mc_text_pt)
        text_run.font.color.rgb = self.styler.colors.text_blue

        # Apply smart RTL formatting based on text content
        self.formatter.apply_smart_rtl_to_paragraph(para, question.question, preserve_alignment=False)
            
    def _render_choices_grid(self, cell, question: MultipleChoiceQuestion, choice_labels: ChoiceLabels) -> None:
        """Render 2x2 choices grid"""
        rows = self.styler.ui.choice_grid.rows
        cols = self.styler.ui.choice_grid.cols
        choices_table = cell.add_table(rows=rows, cols=cols)
        choices_table.autofit = True
        choices_table.alignment = WD_TABLE_ALIGNMENT.CENTER
                
        # Arrange cells right-to-left, top-to-bottom (optimized for 2x2)
        if rows == 2 and cols == 2:
            cell_positions = [
                (0, 1), (0, 0),  # Top row: right, left
                (1, 1), (1, 0)   # Bottom row: right, left
            ]
        else:
            # Generic RTL row-major fallback
            cell_positions = []
            for r in range(rows):
                cols_order = list(range(cols-1, -1, -1))  # right to left
                for c in cols_order:
                    cell_positions.append((r, c))
        
        for i, ((row, col), choice) in enumerate(
            zip(cell_positions, question.choices)
        ):
            choice_cell = choices_table.cell(row, col)

            self._render_single_choice(choice_cell, choice, i, question.answer_key, choice_labels)
    
    def _render_single_choice(self, cell, choice: str, 
                            index: int, answer_key: Optional[int], choice_labels: ChoiceLabels) -> None:
        """Render individual choice option"""
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Choice elements
        if choice_labels.numbering_type == "numbers":
            label_run = para.add_run(f"{choice_labels.numbers[index]}{choice_labels.separator}")
            label_run.font.size = Pt(self.styler.ui.fonts.choice_label_pt)
            label_run.bold = True
            label_run.font.color.rgb = self.styler.colors.text_blue
        else:
            label_run = para.add_run(f"{choice_labels.symbols[index]}{choice_labels.separator}")
            label_run.font.size = Pt(self.styler.ui.fonts.choice_label_pt)
            label_run.bold = True
            label_run.font.color.rgb = self.styler.colors.text_blue
            
        choice_run = para.add_run(self.formatter.format(str(choice)))
        choice_run.font.size = Pt(self.styler.ui.fonts.choice_text_pt)
        choice_run.font.color.rgb = self.styler.colors.text_blue
        
        # Apply smart RTL formatting based on choice text content
        self.formatter.apply_smart_rtl_to_paragraph(para, str(choice), preserve_alignment=False)
        
        # Style based on answer correctness
        if answer_key is not None and index == answer_key:
            self.styler.set_cell_background(cell, self.styler.colors.yellow_highlight)
            choice_run.bold = True
        else:
            self.styler.set_cell_background(cell, self.styler.colors.light_gray)
        
        # self.styler.set_cell_border(cell, self.styler.colors.border_gray, 1)

class EssayRenderer(ContentRenderer):
    """Renders essay questions"""
    
    def render(self, container, question: EssayQuestion, general_labels: GeneralLabels = None) -> None:
        """Render an essay question with answer space"""
        if general_labels is None:
            general_labels = GeneralLabels()
            
        # Main container
        table = container.add_table(rows=1, cols=1)
        table.autofit = True
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        cell = table.cell(0, 0)
        cell._tc.get_or_add_tcPr().append(
            parse_xml(f'<w:tcW w:w="{int(self.styler.ui.dimensions.page_width_in * DXA_PER_INCH)}" w:type="dxa" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')
        )

        # Question header
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        num_run = para.add_run(f" {question.number} ")
        num_run.font.size = Pt(self.styler.ui.fonts.mc_number_pt)
        num_run.bold = True
        num_run.font.color.rgb = self.styler.colors.white

        text_run = para.add_run(f"  {self.formatter.format(question.question)}")
        text_run.font.size = Pt(self.styler.ui.fonts.mc_text_pt)
        text_run.font.color.rgb = self.styler.colors.text_blue

        # Apply smart RTL formatting based on question text
        self.formatter.apply_smart_rtl_to_paragraph(para, question.question, preserve_alignment=False)

        # Show answer if provided
        if question.answer:
            answer_para = cell.add_paragraph()
            answer_run = answer_para.add_run(f"{general_labels.answer_prefix}{self.formatter.format(question.answer)}")
            answer_run.font.size = Pt(self.styler.ui.fonts.essay_answer_pt)
            answer_run.font.color.rgb = self.styler.colors.red_level
            answer_run.bold = True
            # Apply RTL (align right) for answers
            self.formatter.apply_rtl_to_paragraph(answer_para, preserve_alignment=False)
            answer_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Answer space
        for _ in range(self.styler.ui.essay.answer_lines):
            answer_para = cell.add_paragraph()
            answer_para.add_run(general_labels.underline_char * general_labels.underline_count)
            answer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # Apply RTL while preserving CENTER alignment for underlines
            self.formatter.apply_rtl_to_paragraph(answer_para, preserve_alignment=True)

        # Styling
        self.styler.set_cell_background(cell, self.styler.colors.light_gray)
        self.styler.set_cell_border(cell, self.styler.colors.dark_navy, self.styler.ui.borders.width_pt)

class SidebarRenderer(ContentRenderer):
    """Renders sidebar content"""
    
    def render(self, container, sidebar: SidebarConfig, sidebar_labels: SidebarLabels = None) -> None:
        """Render complete sidebar"""
        if sidebar_labels is None:
            sidebar_labels = SidebarLabels()
            
        self._render_section(container, sidebar_labels.before_lesson_title, sidebar.before_lesson)
        self._render_goals(container, sidebar.goals, sidebar_labels)
        self._render_applications(container, sidebar.applications, sidebar_labels)
        self._render_levels(container, sidebar.levels, sidebar_labels)
        self._render_notes(container, sidebar_labels.notes_title)
    
    def _render_section(self, container, title: str, content: str = None) -> None:
        """Render a sidebar section"""
        # Header
        table = container.add_table(rows=1, cols=1)
        table.autofit = True
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.width = Inches(self.styler.ui.dimensions.sidebar_box_width_in)
        
        cell = table.cell(0, 0)
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        run = para.add_run(self.formatter.format(title))
        run.font.size = Pt(self.styler.ui.fonts.header_light_pt)
        run.bold = True
        run.font.color.rgb = self.styler.colors.white
        
        # Apply RTL while preserving CENTER alignment
        self.formatter.apply_rtl_to_paragraph(para, preserve_alignment=True)
        
        self.styler.set_cell_background(cell, self.styler.colors.dark_navy)
        # self.styler.set_cell_border(cell, self.styler.colors.dark_navy)
        
        # Content if provided
        if content:
            content_table = container.add_table(rows=1, cols=1)
            content_table.autofit = True
            content_table.width = Inches(self.styler.ui.dimensions.sidebar_box_width_in)
            
            content_cell = content_table.cell(0, 0)
            content_para = content_cell.paragraphs[0]
            content_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            content_run = content_para.add_run(self.formatter.format(content))
            content_run.font.size = Pt(self.styler.ui.fonts.choice_text_pt)
            content_run.font.color.rgb = self.styler.colors.text_blue
            
            # Apply RTL while preserving CENTER alignment
            self.formatter.apply_rtl_to_paragraph(content_para, preserve_alignment=True)
            
            self.styler.set_cell_background(content_cell, self.styler.colors.light_gray)
            self.styler.set_cell_border(content_cell, self.styler.colors.border_gray)
            
    def _render_goals(self, container, goals: List[str], sidebar_labels: SidebarLabels) -> None:
        """Render lesson objectives"""
        self._render_section(container, sidebar_labels.goals_title)
        
        for goal in goals:
            table = container.add_table(rows=1, cols=1)
            table.autofit = True
            table.width = Inches(self.styler.ui.dimensions.sidebar_box_width_in)
            
            cell = table.cell(0, 0)
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            run = para.add_run(sidebar_labels.goals_bullet + self.formatter.format(goal))
            run.font.size = Pt(self.styler.ui.fonts.choice_text_pt - 1)
            run.font.color.rgb = self.styler.colors.text_blue
            
            # Apply RTL while preserving CENTER alignment
            self.formatter.apply_rtl_to_paragraph(para, preserve_alignment=True)
            
            self.styler.set_cell_background(cell, self.styler.colors.light_gray)
            self.styler.set_cell_border(cell, self.styler.colors.border_gray)
            
    def _render_applications(self, container, applications: List[str], sidebar_labels: SidebarLabels) -> None:
        """Render applications section"""
        self._render_section(container, sidebar_labels.applications_title)
        
        for application in applications:
            table = container.add_table(rows=1, cols=1)
            table.autofit = True
            table.width = Inches(self.styler.ui.dimensions.sidebar_box_width_in)
            
            cell = table.cell(0, 0)
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            run = para.add_run(sidebar_labels.goals_bullet + self.formatter.format(application))
            run.font.size = Pt(self.styler.ui.fonts.choice_text_pt - 1)
            run.font.color.rgb = self.styler.colors.text_blue
            
            # Apply RTL while preserving CENTER alignment
            self.formatter.apply_rtl_to_paragraph(para, preserve_alignment=True)
            
            self.styler.set_cell_background(cell, self.styler.colors.light_gray)
    
    def _render_levels(self, container, levels: List[str], sidebar_labels: SidebarLabels) -> None:
        """Render performance levels"""
        self._render_section(container, sidebar_labels.levels_title)
        
        level_colors = [
            self.styler.colors.green_level,
            self.styler.colors.orange_level,
            self.styler.colors.red_level
        ]
        # Cycle colors if more levels are provided
        color_cycle = cycle(level_colors) if levels and len(levels) > len(level_colors) else None
        
        for idx, level in enumerate(levels):
            color = next(color_cycle) if color_cycle else level_colors[idx]
            table = container.add_table(rows=1, cols=1)
            table.autofit = True
            
            cell = table.cell(0, 0)
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            run = para.add_run(sidebar_labels.levels_bullet + self.formatter.format(level))
            run.font.size = Pt(self.styler.ui.fonts.level_text_pt)
            run.bold = True
            run.font.color.rgb = self.styler.colors.text_blue
            
            # Apply RTL while preserving CENTER alignment
            self.formatter.apply_rtl_to_paragraph(para, preserve_alignment=True)
            
            self.styler.set_cell_background(cell, color)
            self.styler.set_cell_border(cell, color, self.styler.ui.borders.width_pt)
            
    def _render_notes(self, container, title: str) -> None:
        """Render teacher notes section"""
        self._render_section(container, title)
        
        # Notes area
        table = container.add_table(rows=1, cols=1)
        table.autofit = True
        table.width = Inches(self.styler.ui.dimensions.sidebar_box_width_in)
        
        cell = table.cell(0, 0)
        self.styler.set_cell_background(cell, self.styler.colors.light_gray)
        self.styler.set_cell_border(cell, self.styler.colors.border_gray)
        
        for _ in range(3):
            cell.add_paragraph()

class FooterRenderer(ContentRenderer):
    """Renders document footer"""

    def render(self, container, text: str = None) -> None:
        """Render document footer"""
        if not text:
            text = "بسم الله الرحمن الرحيم"

        para = container.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run = para.add_run(self.formatter.format(text))
        run.font.size = Pt(self.styler.ui.fonts.footer_pt)
        run.italic = True
        run.font.color.rgb = self.styler.colors.text_blue

# ============================================================================
# MAIN WORKSHEET GENERATOR
# ============================================================================

class WorksheetGenerator:
    """Main worksheet generator with improved architecture"""
    
    def __init__(self):
        self.colors = Colors()
        self.ui = UIConfig()
        self.formatter = ArabicTextFormatter()
        self.styler = DocumentStyler(self.colors, self.ui)
        self.layout_manager = LayoutManager(self.styler)
        
        # Initialize renderers
        self.header_renderer = HeaderRenderer(self.styler, self.formatter)
        self.section_renderer = SectionHeaderRenderer(self.styler, self.formatter)
        self.mc_renderer = MultipleChoiceRenderer(self.styler, self.formatter)
        self.essay_renderer = EssayRenderer(self.styler, self.formatter)
        self.sidebar_renderer = SidebarRenderer(self.styler, self.formatter)
        # self.footer_renderer = FooterRenderer(self.styler, self.formatter)
    
    def generate(self, data: WorksheetData, filename: str = None) -> str:
        """Generate worksheet document"""
        # Create document
        doc = Document()
        self.styler.setup_document(doc)
        
        # Render header
        self.header_renderer.render(doc, data.title, data.header_config)
        
        # Create layout
        main_container, sidebar_container = self.layout_manager.create_layout(doc, data.layout_type)
        
        # Render content
        self._render_questions(main_container, data)
        
        # Render sidebar if present
        if sidebar_container and data.sidebar:
            self.sidebar_renderer.render(sidebar_container, data.sidebar, data.sidebar_labels)
        
        # Render footer
        # self.footer_renderer.render(main_container)
        
        # Save document
        if not filename:
            timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"worksheet_{timestamp}.docx"
        
        # Iterate through paragraphs or select a specific one
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        p_fmt = paragraph.paragraph_format
                        p_fmt.space_before = Pt(0)  # Adjust as needed
                        p_fmt.space_after = Pt(0)
        doc.save(filename)
        return filename
    
    def _render_questions(self, container, data: WorksheetData) -> None:
        """Render all questions with proper numbering"""
        question_number = 1
        
        # Multiple choice questions
        if data.multiple_choice_questions:
            self.section_renderer.render(container, data.multiple_choice_header, data.general_labels)
            
            for question in data.multiple_choice_questions:
                question.number = question_number
                self.mc_renderer.render(container, question, data.choice_labels)
                question_number += 1
        
        # Essay questions
        if data.essay_questions:
            self.section_renderer.render(container, data.essay_header, data.general_labels)
            
            for question in data.essay_questions:
                question.number = question_number
                self.essay_renderer.render(container, question, data.general_labels)
                question_number += 1

# ============================================================================
# ENHANCED SERVICE INTERFACE
# ============================================================================

class WorksheetService:
    """High-level service interface for worksheet generation"""
    
    def __init__(self):
        self.generator = WorksheetGenerator()
    
    def create_worksheet(self, 
                        data: Union[Dict[str, Any], str], 
                        filename: str = None) -> Dict[str, Any]:
        """
        Create worksheet from data
        
        Args:
            data: Dictionary or JSON string with worksheet data
            filename: Output filename (auto-generated if None)
        
        Returns:
            Result dictionary with status and details
        """
        try:
            # Parse input data
            if isinstance(data, str):
                data = json.loads(data)
            
            # Convert to WorksheetData
            worksheet_data = self._convert_input_data(data)
            
            # Generate worksheet
            output_file = self.generator.generate(worksheet_data, filename)
            
            return {
                "status": "success",
                "message": f"Worksheet successfully created: {output_file}",
                "filename": output_file,
                "layout": worksheet_data.layout_type.value
            }
            
        except json.JSONDecodeError as e:
            return {"status": "error", "message": f"Invalid JSON: {e}"}
        except ValueError as e:
            return {"status": "error", "message": f"Data validation error: {e}"}
        except Exception as e:
            return {"status": "error", "message": f"Generation error: {e}"}
    
    def _convert_input_data(self, data: Dict[str, Any]) -> WorksheetData:
        """Convert input dictionary to WorksheetData object"""
        # Extract basic info
        title = data.get('title', "")
        
        # Convert multiple choice questions
        mc_questions = []
        mc_header = ""
        
        # Convert essay questions
        essay_questions = []
        essay_header = ""
        
        # Handle question_bank format (flat questions list)
        if 'questions' in data and isinstance(data['questions'], list):
            # This is question_bank format with flat questions list
            for q in data['questions']:
                if 'choices' in q and q['choices']:
                    # This is a multiple choice question
                    mc_questions.append(MultipleChoiceQuestion(
                        question=q['question'],
                        choices=q['choices'],
                        answer_key=q.get('answer_key')
                    ))
                else:
                    # This is an essay question
                    essay_questions.append(EssayQuestion(
                        question=q['question'],
                        answer=q.get('answer', '')
                    ))
            
            # Set default headers for question bank
            mc_header = ": اختر الإجابة الصحيحة"
            essay_header = ": أجب عن الأسئلة التالية"
        
        # Handle worksheet format (separate multiple_choice and essay sections)
        if 'multiple_choice' in data and 'questions' in data['multiple_choice']:
            mc_header = data['multiple_choice'].get('header', "")
            for q in data['multiple_choice']['questions']:
                mc_questions.append(MultipleChoiceQuestion(
                    question=q['question'],
                    choices=q['choices'],
                    answer_key=q.get('answer_key')
                ))
        
        if 'essay' in data and 'questions' in data['essay']:
            essay_header = data['essay'].get('header', "")
            for q in data['essay']['questions']:
                essay_questions.append(EssayQuestion(
                    question=q['question'],
                    answer=q.get('answer', '')
                ))
        
        # Convert sidebar if present
        sidebar = None
        if 'sidebar' in data and data['sidebar']:
            sidebar_data = data['sidebar']
            sidebar = SidebarConfig(
                before_lesson=sidebar_data.get('before_lesson', ""),
                goals=sidebar_data.get('goal', sidebar_data.get('goals', [])),
                applications=sidebar_data.get('application', sidebar_data.get('applications', [])),
                levels=sidebar_data.get('level', sidebar_data.get('levels', [])),
                notice=sidebar_data.get('notice', "")
            )
        
        # Extract header configuration
        header_config = HeaderConfig()
        if 'header_config' in data:
            hc_data = data['header_config']
            header_config = HeaderConfig(
                subject_memo=hc_data.get('subject_memo', header_config.subject_memo),
                worksheet_number=hc_data.get('worksheet_number', header_config.worksheet_number),
                name_label=hc_data.get('name_label', header_config.name_label),
                class_label=hc_data.get('class_label', header_config.class_label),
                semester=hc_data.get('semester', header_config.semester),
                grade=hc_data.get('grade', header_config.grade)
            )
        
        # Extract choice labels configuration
        choice_labels = ChoiceLabels()
        if 'choice_labels' in data:
            cl_data = data['choice_labels']
            choice_labels = ChoiceLabels(
                symbols=cl_data.get('symbols', choice_labels.symbols),
                numbers=cl_data.get('numbers', choice_labels.numbers),
                separator=cl_data.get('separator', choice_labels.separator),
                numbering_type=cl_data.get('numbering_type', choice_labels.numbering_type)
            )
        
        # Extract sidebar labels configuration
        sidebar_labels = SidebarLabels()
        if 'sidebar_labels' in data:
            sl_data = data['sidebar_labels']
            sidebar_labels = SidebarLabels(
                before_lesson_title=sl_data.get('before_lesson_title', sidebar_labels.before_lesson_title),
                goals_title=sl_data.get('goals_title', sidebar_labels.goals_title),
                applications_title=sl_data.get('applications_title', sidebar_labels.applications_title),
                levels_title=sl_data.get('levels_title', sidebar_labels.levels_title),
                notes_title=sl_data.get('notes_title', sidebar_labels.notes_title),
                goals_bullet=sl_data.get('goals_bullet', sidebar_labels.goals_bullet),
                levels_bullet=sl_data.get('levels_bullet', sidebar_labels.levels_bullet)
            )
        
        # Extract general labels configuration
        general_labels = GeneralLabels()
        if 'general_labels' in data:
            gl_data = data['general_labels']
            general_labels = GeneralLabels(
                section_icon=gl_data.get('section_icon', general_labels.section_icon),
                answer_prefix=gl_data.get('answer_prefix', general_labels.answer_prefix),
                underline_char=gl_data.get('underline_char', general_labels.underline_char),
                underline_count=gl_data.get('underline_count', general_labels.underline_count)
            )
        
        return WorksheetData(
            title=title,
            multiple_choice_header=mc_header,
            essay_header=essay_header,
            multiple_choice_questions=mc_questions,
            essay_questions=essay_questions,
            sidebar=sidebar,
            header_config=header_config,
            choice_labels=choice_labels,
            sidebar_labels=sidebar_labels,
            general_labels=general_labels
        )
    
    def validate_data(self, data: Union[Dict[str, Any], str]) -> Dict[str, Any]:
        """Validate worksheet data structure"""
        try:
            if isinstance(data, str):
                data = json.loads(data)
            
            errors = []
            warnings = []
            
            # Check for questions
            has_mc = 'multiple_choice' in data and 'questions' in data['multiple_choice']
            has_essay = 'essay' in data and 'questions' in data['essay']
            
            if not has_mc and not has_essay:
                errors.append("At least one question type must be provided")
            
            # Validate multiple choice
            if has_mc:
                for i, q in enumerate(data['multiple_choice']['questions']):
                    if 'question' not in q:
                        errors.append(f"MC question {i+1}: missing 'question' field")
                    if 'choices' not in q or len(q['choices']) != 4:
                        errors.append(f"MC question {i+1}: must have exactly 4 choices")
                    if 'answer_key' in q and not (0 <= q['answer_key'] < 4):
                        errors.append(f"MC question {i+1}: answer_key must be 0-3")
            
            # Validate essay questions
            if has_essay:
                for i, q in enumerate(data['essay']['questions']):
                    if 'question' not in q:
                        errors.append(f"Essay question {i+1}: missing 'question' field")
            
            # Check sidebar
            if 'sidebar' not in data or not data['sidebar']:
                warnings.append("No sidebar data provided - will use full-width layout")
            
            return {
                "valid": len(errors) == 0,
                "errors": errors,
                "warnings": warnings
            }
            
        except Exception as e:
            return {
                "valid": False,
                "errors": [f"Validation error: {e}"],
                "warnings": []
            }

# ============================================================================
# CONVENIENCE FUNCTIONS
# ============================================================================

def create_worksheet_from_json(json_data: str, filename: str = None) -> str:
    """Convenience function to create worksheet from JSON"""
    service = WorksheetService()
    result = service.create_worksheet(json_data, filename)
    if result["status"] == "success":
        return result["filename"]
    else:
        raise Exception(result["message"])

def create_worksheet_from_dict(data: Dict[str, Any], filename: str = None) -> str:
    """Convenience function to create worksheet from dictionary"""
    service = WorksheetService()
    result = service.create_worksheet(data, filename)
    if result["status"] == "success":
        return result["filename"]
    else:
        raise Exception(result["message"])


if __name__ == "__main__":
    # Demo usage
    sample_data = {
        "title": "ورقة عمل الرياضيات",
        "multiple_choice": {
            "header": ": اختر الإجابة الصحيحة",
            "questions": [
                {
                    "question": "كم يساوي ٥ + ٣؟",
                    "choices": ["٦", "٧", "٨", "٩"],
                    "answer_key": 2
                },
                {
                    "question": "كم يساوي ٥ + ٣؟",
                    "choices": ["٦", "٧", "٨", "٩"],
                    "answer_key": 2
                },
                {
                    "question": "كم يساوي ٥ + ٣؟",
                    "choices": ["٦", "٧", "٨", "٩"],
                    "answer_key": 2
                },{
                    "question": "كم يساوي ٥ + ٣؟",
                    "choices": ["٦", "٧", "٨", "٩"],
                    "answer_key": 2
                },
            ]
        },
        "essay": {
            "header": ": أجب عن الأسئلة التالية",
            "questions": [
                {
                    "question": "اشرح أهمية الرياضيات في الحياة اليومية",
                    "answer": "الرياضيات مهمة للحسابات والتفكير المنطقي"
                },
                {
                    "question": "كيف يمكن استخدام الرياضيات في الحياة اليومية؟",
                    "answer": "يمكن استخدام الرياضيات في التسوق، الطبخ، وإدارة الوقت."
                },{
                    "question": "كيف يمكن استخدام الرياضيات في الحياة اليومية؟",
                    "answer": "يمكن استخدام الرياضيات في التسوق، الطبخ، وإدارة الوقت."
                },
            ]
        },
        "sidebar": {
            "before_lesson": "مقدمة في الحساب",
            "goal": ["تعلم الجمع", "فهم الأرقام"],
            "application": ["التسوق", "الطبخ"],
            "level": ["ممتاز", "متوسط", "منخفض"],
            "notice": "راجع الطلاب الضعاف"
        },
        "header_config": {
            "subject_memo": "مذكرة الرياضيات",
            "worksheet_number": "ورقة عمل 1",
            "name_label": " :الاسم",
            "class_label": " :الصف",
            "semester": "الفصل الأول",
            "grade": "الصف الثالث"
        },
        "choice_labels": {
            "symbols": ["أ", "ب", "ج", "د"],
            "numbers": ["١", "٢", "٣", "٤"],
            "separator": " - ",
            "numbering_type": "numbers"
        },
        "sidebar_labels": {
            "before_lesson_title": "ما قبل الدرس",
            "goals_title": "أهداف الدرس",
            "applications_title": "تطبيقات الدرس",
            "levels_title": "مستويات الاتقان",
            "notes_title": "ملاحظات المعلم",
            "goals_bullet": "✓ ",
            "levels_bullet": "〇 \t"
        },
        "general_labels": {
            "section_icon": "",
            "answer_prefix": ":الإجابة ",
            "underline_char": "_",
            "underline_count": 100
        }
    }
    
    # Create worksheet
    service = WorksheetService()
    result = service.create_worksheet(sample_data, "demo_worksheet.docx")
    print(f"Result: {result}")
